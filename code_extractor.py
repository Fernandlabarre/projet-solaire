import os
from docx import Document
from docx.shared import Pt

def add_file_content_to_doc(document, file_path):
    document.add_heading(file_path, level=2)
    try:
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            content = f.read()
        paragraph = document.add_paragraph()
        run = paragraph.add_run(content)
        font = run.font
        font.name = 'Courier New'
        font.size = Pt(10)
    except Exception as e:
        print(f"Impossible de lire {file_path} : {e}")



def main():
    # Liste des chemins de base que l'on veut explorer
    base_paths = [
        "C:/Users/rescue123/Documents/DEV/CRM/backend",
                "C:/Users/rescue123/Documents/DEV/CRM/crm-frontend",

    ]

    output_file = "extractor.docx"

    document = Document()
    document.add_heading('Export de code', level=1)

    # On parcourt chacun des répertoires de base
    for base_path in base_paths:
        for root, dirs, files in os.walk(base_path):
            # On retire node_modules et .git des dossiers à explorer
            if 'node_modules' in dirs:
                dirs.remove('node_modules')
            if '.git' in dirs:
                dirs.remove('.git')
            if 'dist' in dirs:
                dirs.remove('dist')
            for file in files:
                # Ignorer package-lock.json
                if file == 'package-lock.json':
                    continue

                file_path = os.path.join(root, file)
                add_file_content_to_doc(document, file_path)

    document.save(output_file)
    print(f"Le fichier {output_file} a été généré avec succès.")

if __name__ == "__main__":
    main()
